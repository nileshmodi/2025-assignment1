from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os

def format_title_and_toc(input_path, output_path):
    doc = Document(input_path)
    new_doc = Document()

    # --- Title Page ---
    section = new_doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Title (assume first paragraph is the title)
    title = doc.paragraphs[0].text.strip()
    p = new_doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph('')

    # Author/Version/Date (if present in next few paragraphs)
    for i in range(1, 5):
        if i < len(doc.paragraphs):
            text = doc.paragraphs[i].text.strip()
            if text:
                p = new_doc.add_paragraph()
                run = p.add_run(text)
                run.font.size = Pt(12)
                run.font.italic = True
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph('')
    new_doc.add_paragraph('')

    # Horizontal line for separation
    p = new_doc.add_paragraph()
    run = p.add_run('_' * 60)
    run.font.color.rgb = RGBColor(180, 180, 180)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    new_doc.add_paragraph('')

    # --- Table of Contents ---
    toc_title = new_doc.add_paragraph()
    run = toc_title.add_run('Table of Contents')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    toc_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    new_doc.add_paragraph('')

    # Copy over the TOC (assume it's in the next few paragraphs)
    for para in doc.paragraphs[5:15]:
        text = para.text.strip()
        if text:
            p = new_doc.add_paragraph(text)
            p.style = 'List Number'

    # Save the new document
    new_doc.save(output_path)

if __name__ == "__main__":
    format_title_and_toc("NILESH.docx", "NILESH_V2.docx")
