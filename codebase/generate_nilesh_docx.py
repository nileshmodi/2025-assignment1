import os
import re
import requests
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_markdown_table(doc, md_table):
    lines = [line for line in md_table.strip().split('\n') if line.strip()]
    if len(lines) < 2:
        return
    headers = [h.strip() for h in lines[0].strip('|').split('|')]
    rows = [
        [cell.strip() for cell in row.strip('|').split('|')]
        for row in lines[2:]
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for row in rows:
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = cell
    doc.add_paragraph('')

def render_mermaid_to_img(mermaid_code, img_path):
    url = 'https://kroki.io/mermaid/png'
    resp = requests.post(url, data=mermaid_code.encode('utf-8'))
    if resp.status_code == 200:
        with open(img_path, 'wb') as f:
            f.write(resp.content)
        return True
    return False

def add_mermaid_diagram(doc, mermaid_code, diagram_name):
    img_path = f"{diagram_name}.png"
    if render_mermaid_to_img(mermaid_code, img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        os.remove(img_path)
    else:
        doc.add_paragraph(f"[Diagram could not be rendered: {diagram_name}]")

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_paragraph(doc, text):
    doc.add_paragraph(text)

def parse_and_write_docx(md_path, docx_path):
    with open(md_path, encoding='utf-8') as f:
        md = f.read()
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    code_block_re = re.compile(r'```(\w+)?\n([\s\S]*?)```', re.MULTILINE)
    pos = 0
    for m in code_block_re.finditer(md):
        before = md[pos:m.start()]
        code_type = m.group(1)
        code_content = m.group(2)
        process_markdown_segment(doc, before)
        if code_type == 'mermaid':
            add_mermaid_diagram(doc, code_content, f"diagram_{m.start()}")
        else:
            doc.add_paragraph(code_content, style='Intense Quote')
        pos = m.end()
    process_markdown_segment(doc, md[pos:])
    doc.save(docx_path)

def process_markdown_segment(doc, text):
    lines = text.split('\n')
    buffer = []
    in_table = False
    table_lines = []
    for line in lines:
        if re.match(r'^\s*\|.*\|\s*$', line):
            in_table = True
            table_lines.append(line)
        else:
            if in_table:
                add_markdown_table(doc, '\n'.join(table_lines))
                table_lines = []
                in_table = False
            if line.strip().startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                add_heading(doc, line.strip('# ').strip(), level=level)
            elif line.strip():
                buffer.append(line)
            elif buffer:
                add_paragraph(doc, '\n'.join(buffer).strip())
                buffer = []
    if in_table:
        add_markdown_table(doc, '\n'.join(table_lines))
    if buffer:
        add_paragraph(doc, '\n'.join(buffer).strip())

if __name__ == '__main__':
    parse_and_write_docx('psd.md', 'NILESH.docx')
