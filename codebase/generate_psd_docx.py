from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

# Formatting constants
MARGIN = Cm(2.54)
FONT_FAMILY = 'Calibri'
BODY_SIZE = 11
H1_SIZE = 18
H2_SIZE = 14
H3_SIZE = 12
TABLE_SIZE = 10

# Helper functions
def set_margins(section):
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN

def set_header_footer(doc, title, version_date):
    section = doc.sections[0]
    header = section.header
    footer = section.footer
    # Header: Title (left), version/date (right)
    table = header.add_table(rows=1, cols=2, width=Cm(15.92))
    table.allow_autofit = True
    table.cell(0, 0).text = title
    table.cell(0, 1).text = version_date
    for i, cell in enumerate(table.rows[0].cells):
        for p in cell.paragraphs:
            p.runs[0].font.size = Pt(10)
            p.runs[0].font.name = FONT_FAMILY
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.RIGHT
    # Footer: Page number centered
    p = footer.paragraphs[0]
    p.text = "Page "
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    p._element.append(fldChar)
    p._element.append(instrText)
    p._element.append(fldChar2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.name = FONT_FAMILY

def add_heading(doc, text, level):
    if level == 1:
        p = doc.add_heading(text, level=1)
        run = p.runs[0]
        run.font.size = Pt(H1_SIZE)
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif level == 2:
        p = doc.add_heading(text, level=2)
        run = p.runs[0]
        run.font.size = Pt(H2_SIZE)
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif level == 3:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(H3_SIZE)
        run.font.bold = True
        run.font.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_table(doc, data, header=True, banded=True):
    rows = len(data)
    cols = len(data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(TABLE_SIZE)
                    run.font.name = FONT_FAMILY
                if header and i == 0:
                    for run in p.runs:
                        run.font.bold = True
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if banded and i % 2 == 1 and i != 0:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'F2F2F2')
                tcPr.append(shd)
    return table

def add_paragraph(doc, text, italic=False, bold=False, size=BODY_SIZE, justify=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = FONT_FAMILY
    run.font.italic = italic
    run.font.bold = bold
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def main():
    doc = Document()
    section = doc.sections[0]
    set_margins(section)
    set_header_footer(doc, "Calorie Tracker PSD", f"v1.0 / {datetime.now().strftime('%Y-%m-%d')}")

    # 1. Introduction
    add_heading(doc, "1. Introduction", 1)
    add_paragraph(doc, "This document specifies the requirements and design for the Calorie Tracker web application. It is intended for personal use, enabling users to track their daily calorie intake and visualize their dietary habits. The audience includes developers, testers, and stakeholders interested in the product's features and architecture.")
    add_paragraph(doc, "Background: Calorie tracking is a common need for individuals seeking to manage their diet and health. This product provides a simple, web-based solution.")
    add_paragraph(doc, "Audience: End users (individuals), developers, and testers.")
    add_paragraph(doc, "Definitions / Acronyms: JWT: JSON Web Token; CRUD: Create, Read, Update, Delete")

    # 2. Product Overview
    add_heading(doc, "2. Product Overview", 1)
    add_paragraph(doc, "A web-based calorie tracking application allowing users to log, view, and analyze their daily calorie intake.")
    add_paragraph(doc, "Who is it for? Any individual interested in monitoring their diet and calorie consumption.")
    add_paragraph(doc, "Key features at a glance: Google-based authentication; Add, edit, and delete meal entries with calorie values; View calorie intake by day and over time (1, 2, or 4 weeks); Visualize data with bar charts; Generate test data for demo/testing; Responsive, modern UI with light/dark theme toggle.")

    # 3. Goals & Success Metrics
    add_heading(doc, "3. Goals & Success Metrics", 1)
    add_paragraph(doc, "Goal 1: Enable users to easily log and review their calorie intake. Metric: Number of active users logging meals per week.")
    add_paragraph(doc, "Goal 2: Provide clear visual feedback on dietary habits. Metric: User engagement with stats/charts (page views, time spent).")
    add_paragraph(doc, "Goal 3: Ensure secure, reliable access to personal data. Metric: Zero unauthorized access incidents; uptime >99%.")

    # 4. Scope
    add_heading(doc, "4. Scope", 1)
    add_heading(doc, "4.1 In-Scope", 2)
    add_paragraph(doc, "User authentication via Google; CRUD operations for calorie entries (add, edit, delete, view); Daily and historical calorie statistics; Data visualization (bar charts); Test data generation for user accounts; Light/dark theme support.")
    add_heading(doc, "4.2 Out-of-Scope", 2)
    add_paragraph(doc, "Integration with wearables or external health apps; Social features (sharing, friends); Nutrition advice or recommendations; Multi-user roles (admin, nutritionist, etc.).")

    # 5. User Personas
    add_heading(doc, "5. User Personas", 1)
    add_table(doc, [["Persona", "Description", "Goals/Needs"], ["User", "Any individual tracking their own calories", "Log meals, view stats, manage data"]])

    # 6. User Stories & Use Cases
    add_heading(doc, "6. User Stories & Use Cases", 1)
    user_stories = [
        "As a user, I want to sign in with my Google account, so that my data is secure and personalized.",
        "As a user, I want to add a meal with a description and calorie value, so that I can track what I eat.",
        "As a user, I want to edit or delete a meal entry, so that I can correct mistakes.",
        "As a user, I want to view my calorie intake for today and previous days, so that I can monitor my diet.",
        "As a user, I want to see a bar chart of my daily calorie intake over the past weeks, so that I can spot trends.",
        "As a user, I want to generate test data, so that I can see how the app works with sample entries.",
        "As a user, I want to switch between light and dark themes, so that the UI is comfortable for me.",
        "As a user, I want to be logged out if my session expires or is invalid, so that my data remains secure.",
        "As a user, I want to be notified if login fails or if I try to access a page without being authenticated."
    ]
    for i, story in enumerate(user_stories, 1):
        add_paragraph(doc, f"{i}. {story}")

    # 7. Functional Requirements
    add_heading(doc, "7. Functional Requirements", 1)
    fr_table = [
        ["ID", "Requirement"],
        ["FR-1", "The system shall allow users to sign in using Google OAuth."],
        ["FR-2", "The system shall allow users to add a meal entry with description and calorie value."],
        ["FR-3", "The system shall allow users to edit or delete their own meal entries."],
        ["FR-4", "The system shall display a list of meal entries for the authenticated user."],
        ["FR-5", "The system shall provide daily and historical calorie statistics (1, 2, 4 weeks)."],
        ["FR-6", "The system shall visualize calorie data using bar charts."],
        ["FR-7", "The system shall allow users to generate test data (mock entries for 30 days)."],
        ["FR-8", "The system shall support light and dark UI themes."],
        ["FR-9", "The system shall log out users and restrict access if authentication fails or expires."],
        ["FR-10", "The system shall validate all user input (e.g., required fields, numeric calories > 0)."],
    ]
    add_table(doc, fr_table)

    # 8. Non-Functional Requirements
    add_heading(doc, "8. Non-Functional Requirements", 1)
    nfr_table = [
        ["Category", "Requirement"],
        ["Performance", "Should support at least 100 concurrent users with <2s response time."],
        ["Security", "All API endpoints require authentication (JWT)."],
        ["Security", "User data is isolated; users cannot access others' data."],
        ["Security", "Uses Google OAuth for authentication."],
        ["Privacy", "No personal data is shared externally."],
        ["Availability", "99% uptime (Dockerized, restart policies in place)."],
        ["Scalability", "Can be scaled horizontally via Docker containers."],
        ["Usability", "Responsive design, accessible to anyone on modern browsers."],
        ["Accessibility", "Color contrast and theme options for visual comfort."],
    ]
    add_table(doc, nfr_table)

    # 9. System Architecture
    add_heading(doc, "9. System Architecture", 1)
    add_paragraph(doc, "Overview Diagram: [Diagram: System Architecture (to be generated)]")
    add_paragraph(doc, "Components: Frontend: React (Vite, TypeScript); Backend: Node.js (NestJS, TypeORM, SQLite); Database: SQLite (local file); External services: Google OAuth, Mock image analysis API; Dockerized for local development and deployment.")

    # 10. Data Model
    add_heading(doc, "10. Data Model", 1)
    add_paragraph(doc, "Entity Relationship Diagram: [Diagram: ERD (to be generated)]")
    dm_table = [
        ["Entity", "Attributes"],
        ["User", "id, email, name, picture, accessToken, status, emailVerified, createdAt, updatedAt"],
        ["Calorie", "id, description, calories, userId, createdAt, updatedAt, deleted"],
    ]
    add_table(doc, dm_table)

    # 11. UI/UX Mockups (optional)
    add_heading(doc, "11. UI/UX Mockups (optional)", 1)
    add_paragraph(doc, "The UI is modern, responsive, and supports both light and dark themes. Main screens: Landing (login), Stats (dashboard with chart and table), Modal dialogs for add/edit. Branding: Custom logo, accessible color palette.")

    # 12. Acceptance Criteria
    add_heading(doc, "12. Acceptance Criteria", 1)
    add_paragraph(doc, "FR-1: Given a new user, when they sign in with Google, their account is created and they are authenticated. Negative case: Invalid or expired token shows an error and does not log in.")
    add_paragraph(doc, "FR-2: Given an authenticated user, when they submit a valid meal entry, it is added to their log. Negative case: Missing or invalid fields show an error.")
    add_paragraph(doc, "FR-3: Given an authenticated user, when they edit or delete a meal, the change is reflected in their log. Negative case: Attempting to edit/delete another user's entry is denied.")
    add_paragraph(doc, "FR-4: Given an authenticated user, when they visit the stats page, they see their own meal entries and stats. Negative case: Unauthenticated users are redirected to login.")

    # 13. Constraints & Assumptions
    add_heading(doc, "13. Constraints & Assumptions", 1)
    add_paragraph(doc, "Constraints: Only Google OAuth is supported for authentication. Data is stored locally in SQLite (no cloud DB). No integrations with external health/nutrition APIs.")
    add_paragraph(doc, "Assumptions: Users have access to a modern web browser. Users have an active internet connection.")

    # 14. Glossary
    add_heading(doc, "14. Glossary", 1)
    glossary_table = [
        ["Term", "Definition"],
        ["JWT", "JSON Web Token"],
        ["CRUD", "Create, Read, Update, Delete"],
        ["OAuth", "Open Authorization (Google login)"],
        ["ERD", "Entity Relationship Diagram"],
        ["UI/UX", "User Interface / User Experience"],
    ]
    add_table(doc, glossary_table)

    # 15. Appendices
    add_heading(doc, "15. Appendices", 1)
    add_paragraph(doc, "A. API Endpoints Listing: POST /api/auth/token-signin — Google login; GET /api/calories — List meal entries; POST /api/calories — Add meal entry; PUT /api/calories/:id — Edit meal entry; DELETE /api/calories/:id — Delete meal entry; GET /api/calories/by-day — Daily calorie stats; POST /api/calories/test-data — Generate test data.")
    add_paragraph(doc, "B. Sequence Diagrams: [Diagram: Login Sequence (to be generated)]; [Diagram: Add/Edit/Delete Meal (to be generated)]")
    add_paragraph(doc, "C. Deployment Guide: See README.md for Docker and manual setup instructions.")

    doc.save("PDS_2239_Nilesh_Modi.docx")

if __name__ == "__main__":
    main()