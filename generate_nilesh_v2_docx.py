import os
import re
import requests
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE

def add_title_page(doc, title, version, author, date, logo_path=None):
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_height = 16840  # A4
    section.page_width = 11907
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    if logo_path and os.path.exists(logo_path):
        doc.add_picture(logo_path, width=Inches(1.5))
        doc.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = doc.add_paragraph(f"Version: {version}")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = doc.add_paragraph(f"Author: {author}")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = doc.add_paragraph(f"Date: {date}")
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()

def add_toc(doc):
    p = doc.add_paragraph('Table of Contents')
    p.style = 'Heading 1'
    run = p.add_run('\n(This TOC will update in Word: References > Update Table)')
    run.font.size = Pt(9)
    doc.add_paragraph('')
    # Insert a TOC field (Word will update it)
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'TOC \o "1-3" \h \z \u')
    doc._body._element.append(fldSimple)
    doc.add_page_break()

def add_markdown_table(doc, md_table):
    lines = [line for line in md_table.strip().split('\n') if line.strip()]
    if len(lines) < 2:
        return
    headers = [h.strip() for h in lines[0].strip('|').split('|')]
    rows = [
        [cell.strip() for cell in row.strip('|').split('|')]
        for row in lines[2:]
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
    for idx, row in enumerate(rows):
        row_cells = table.add_row().cells
        for i, cell in enumerate(row):
            row_cells[i].text = cell
            for p in row_cells[i].paragraphs:
                for run in p.runs:
                    if idx % 2 == 0:
                        run.font.color.rgb = RGBColor(0, 0, 0)
        if idx % 2 == 1:
            for cell in row_cells:
                cell.shading_elm = OxmlElement('w:shd')
                cell.shading_elm.set(qn('w:fill'), 'E6F0FA')
                cell._tc.get_or_add_tcPr().append(cell.shading_elm)
    doc.add_paragraph('')

def render_mermaid_to_img(mermaid_code, img_path):
    url = 'https://kroki.io/mermaid/png'
    resp = requests.post(url, data=mermaid_code.encode('utf-8'))
    if resp.status_code == 200:
        with open(img_path, 'wb') as f:
            f.write(resp.content)
        return True
    return False

def add_mermaid_diagram(doc, mermaid_code, diagram_name, caption=None):
    img_path = f"{diagram_name}.png"
    if render_mermaid_to_img(mermaid_code, img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cap.runs[0].font.italic = True
        os.remove(img_path)
    else:
        doc.add_paragraph(f"[Diagram could not be rendered: {diagram_name}]")

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    elif level == 2:
        p.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    else:
        p.runs[0].font.color.rgb = RGBColor(51, 102, 153)

def add_paragraph(doc, text, style=None, italic=False, note=False):
    p = doc.add_paragraph(text, style=style)
    if note:
        for run in p.runs:
            run.font.italic = True
            run.font.color.rgb = RGBColor(102, 102, 102)
    if italic:
        for run in p.runs:
            run.font.italic = True
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

def add_code_block(doc, code):
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(51, 51, 51)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.1
    p.style = 'No Spacing'
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'F2F2F2')
    p._p.get_or_add_pPr().append(shading_elm)

def process_markdown_segment(doc, text):
    lines = text.split('\n')
    buffer = []
    in_table = False
    table_lines = []
    for line in lines:
        if re.match(r'^\s*\|.*\|\s*$', line):
            in_table = True
            table_lines.append(line)
        else:
            if in_table:
                add_markdown_table(doc, '\n'.join(table_lines))
                table_lines = []
                in_table = False
            if line.strip().startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                add_heading(doc, line.strip('# ').strip(), level=level)
            elif line.strip().startswith('**Note:**'):
                add_paragraph(doc, line.strip(), note=True)
            elif line.strip():
                buffer.append(line)
            elif buffer:
                add_paragraph(doc, '\n'.join(buffer).strip())
                buffer = []
    if in_table:
        add_markdown_table(doc, '\n'.join(table_lines))
    if buffer:
        add_paragraph(doc, '\n'.join(buffer).strip())

def parse_and_write_docx(md_path, docx_path):
    with open(md_path, encoding='utf-8') as f:
        md = f.read()
    doc = Document()
    # Title page
    add_title_page(doc, 'Calorie Tracker - Product Specification Document', '1.0', 'Nilesh Modi', '2025-07-15')
    # TOC
    add_toc(doc)
    # Split by code blocks (```...```), keep track of diagrams
    code_block_re = re.compile(r'```(\w+)?\n([\s\S]*?)```', re.MULTILINE)
    pos = 0
    for m in code_block_re.finditer(md):
        before = md[pos:m.start()]
        code_type = m.group(1)
        code_content = m.group(2)
        process_markdown_segment(doc, before)
        if code_type == 'mermaid':
            add_mermaid_diagram(doc, code_content, f"diagram_{m.start()}", caption='System Diagram')
        else:
            add_code_block(doc, code_content)
        pos = m.end()
    process_markdown_segment(doc, md[pos:])
    doc.save(docx_path)

if __name__ == '__main__':
    parse_and_write_docx('psd.md', 'NILESH_V2.docx')
